import customtkinter as ctk
from tkinter import messagebox
import os
from docx import Document
import docx.shared  # Add this explicit import for docx.shared
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from PIL import Image
import subprocess  # For opening files after creation
import platform    # For detecting operating system

# Add this import at the top of the file with other imports
try:
    from file_utils import open_generated_file
except ImportError:
    # Fallback implementation if the module is not found
    def open_generated_file(filepath):
        try:
            filepath = os.path.abspath(filepath)
            if platform.system() == 'Windows':
                os.startfile(filepath)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.Popen(['open', filepath])
            else:  # Linux
                subprocess.Popen(['xdg-open', filepath])
            return True
        except Exception as e:
            print(f"Error opening file: {e}")
            messagebox.showwarning(
                "Warning", 
                f"Could not open the file automatically.\nThe file has been saved to:\n{filepath}"
            )
            return False

class ResumeBuilder:
    def __init__(self):
        self.app = ctk.CTk()
        self.app.title("Modern Resume Builder")
        self.app.geometry("1000x700")
        self.app.resizable(True, True)
        
        # Set appearance
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")
        
        # Create main container
        self.main_container = ctk.CTkFrame(self.app, fg_color="transparent")
        self.main_container.pack(fill="both", expand=True, padx=15, pady=15)
        
        # Create sidebar and content area
        self.create_sidebar()
        self.create_content_area()
        
        # Initialize with the form view
        self.show_form()
        
    def create_sidebar(self):
        # Load resources directory
        resources_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "resources")
        os.makedirs(resources_dir, exist_ok=True)
        
        # Sidebar frame
        self.sidebar = ctk.CTkFrame(self.main_container, width=220, corner_radius=10)
        self.sidebar.pack(side="left", fill="y", padx=(0, 15))
        self.sidebar.pack_propagate(False)
        
        # App title
        title_frame = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        title_frame.pack(fill="x", padx=20, pady=(30, 20))
        
        app_title = ctk.CTkLabel(
            title_frame, 
            text="RESUME BUILDER", 
            font=ctk.CTkFont(family="Arial", size=20, weight="bold")
        )
        app_title.pack(anchor="w")
        
        tagline = ctk.CTkLabel(
            title_frame,
            text="Create professional resumes",
            font=ctk.CTkFont(family="Arial", size=12),
            text_color=("gray70", "gray50")
        )
        tagline.pack(anchor="w")
        
        # Navigation menu
        self.nav_frame = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        self.nav_frame.pack(fill="x", pady=20)
        
        # Navigation buttons
        self.nav_buttons = {}
        
        self.nav_buttons["form"] = self.create_nav_button(
            "📝 Create Resume", 
            self.show_form
        )
        
        self.nav_buttons["preview"] = self.create_nav_button(
            "👁️ Preview", 
            self.show_preview
        )
        
        self.nav_buttons["export"] = self.create_nav_button(
            "📤 Export Options", 
            self.show_export
        )
        
        self.nav_buttons["help"] = self.create_nav_button(
            "❓ Help & Tips", 
            self.show_help
        )
        
        # Current active view
        self.active_view = "form"
        self.update_nav_highlight()
        
        # Footer
        footer_frame = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        footer_frame.pack(side="bottom", fill="x", padx=20, pady=20)
        
        version_label = ctk.CTkLabel(
            footer_frame,
            text="Version 1.0",
            font=ctk.CTkFont(size=11),
            text_color=("gray70", "gray50")
        )
        version_label.pack(side="left")
    
    def create_nav_button(self, text, command):
        button = ctk.CTkButton(
            self.nav_frame,
            text=text,
            anchor="w",
            font=ctk.CTkFont(size=14),
            height=40,
            corner_radius=8,
            fg_color="transparent",
            text_color=("gray10", "gray90"),
            hover_color=("gray80", "gray30"),
            command=command
        )
        button.pack(fill="x", padx=10, pady=5)
        return button
    
    def update_nav_highlight(self):
        for view, button in self.nav_buttons.items():
            if view == self.active_view:
                button.configure(fg_color=("gray80", "gray30"))
            else:
                button.configure(fg_color="transparent")
    
    def create_content_area(self):
        # Main content area
        self.content_area = ctk.CTkFrame(self.main_container, corner_radius=10)
        self.content_area.pack(side="right", fill="both", expand=True)
        
        # Create different content frames but don't show them yet
        self.form_frame = ctk.CTkFrame(self.content_area, fg_color="transparent")
        self.preview_frame = ctk.CTkFrame(self.content_area, fg_color="transparent")
        self.export_frame = ctk.CTkFrame(self.content_area, fg_color="transparent")
        self.help_frame = ctk.CTkFrame(self.content_area, fg_color="transparent")
        
        # Initialize form elements
        self.create_form_elements()
        self.create_preview_elements()
        self.create_export_elements()
        self.create_help_elements()
    
    def create_form_elements(self):
        # Initialize entry_widgets dictionary first
        self.entry_widgets = {}
        
        # Form header
        header_frame = ctk.CTkFrame(self.form_frame, fg_color="transparent")
        header_frame.pack(fill="x", padx=20, pady=(20, 10))
        
        header_title = ctk.CTkLabel(
            header_frame,
            text="Create Your Resume",
            font=ctk.CTkFont(family="Arial", size=24, weight="bold")
        )
        header_title.pack(side="left")
        
        clear_button = ctk.CTkButton(
            header_frame,
            text="Clear Form",
            font=ctk.CTkFont(size=13),
            width=110,
            height=32,
            corner_radius=6,
            command=self.clear_form
        )
        clear_button.pack(side="right")
        
        # Scrollable form container
        form_container = ctk.CTkScrollableFrame(self.form_frame, fg_color="transparent")
        form_container.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        
        # Progress indicator
        progress_frame = ctk.CTkFrame(form_container, fg_color="transparent")
        progress_frame.pack(fill="x", pady=(0, 20))
        
        self.progress_bar = ctk.CTkProgressBar(progress_frame, height=10, corner_radius=5)
        self.progress_bar.pack(fill="x")
        self.progress_bar.set(0)
        
        progress_text = ctk.CTkLabel(
            progress_frame,
            text="Fill out your information below",
            font=ctk.CTkFont(size=12),
            text_color=("gray70", "gray50")
        )
        progress_text.pack(anchor="e", pady=(5, 0))
        
        # Personal Information Section
        self.create_section(
            form_container, 
            "Personal Information", 
            "👤", 
            [
                {"label": "Full Name", "var_name": "name", "required": True},
                {"label": "Email", "var_name": "email", "required": True},
                {"label": "Phone Number", "var_name": "phone", "required": False},
                {"label": "Address", "var_name": "address", "required": False}
            ]
        )
        
        # Professional Summary Section
        summary_section = self.create_fancy_section(form_container, "Professional Summary", "📋")
        self.summary_text = self.create_textbox(summary_section, height=100)
        self.summary_text.pack(fill="x", padx=15, pady=(0, 15))
        
        # Skills Section
        skills_section = self.create_fancy_section(form_container, "Skills", "🔧")
        hint_label = ctk.CTkLabel(
            skills_section,
            text="Enter skills separated by commas (e.g., Python, JavaScript, Project Management)",
            font=ctk.CTkFont(size=12),
            text_color=("gray70", "gray50")
        )
        hint_label.pack(anchor="w", padx=15, pady=(0, 5))
        
        self.skills_text = self.create_textbox(skills_section, height=80)
        self.skills_text.pack(fill="x", padx=15, pady=(0, 15))
        
        # Education Section
        education_section = self.create_fancy_section(form_container, "Education", "🎓")
        hint_label = ctk.CTkLabel(
            education_section,
            text="Include degree, institution, location and graduation year",
            font=ctk.CTkFont(size=12),
            text_color=("gray70", "gray50")
        )
        hint_label.pack(anchor="w", padx=15, pady=(0, 5))
        
        self.education_text = self.create_textbox(education_section, height=100)
        self.education_text.pack(fill="x", padx=15, pady=(0, 15))
        
        # Experience Section
        experience_section = self.create_fancy_section(form_container, "Work Experience", "💼")
        hint_label = ctk.CTkLabel(
            experience_section,
            text="Include job title, company, location, dates and achievements",
            font=ctk.CTkFont(size=12),
            text_color=("gray70", "gray50")
        )
        hint_label.pack(anchor="w", padx=15, pady=(0, 5))
        
        self.experience_text = self.create_textbox(experience_section, height=120)
        self.experience_text.pack(fill="x", padx=15, pady=(0, 15))
        
        # Update progress bar when fields change
        self.bind_progress_updates()
    
    def create_section(self, parent, title, icon, fields):
        section = self.create_fancy_section(parent, title, icon)
        
        # Create fields grid
        fields_frame = ctk.CTkFrame(section, fg_color="transparent")
        fields_frame.pack(fill="x", padx=15, pady=(0, 15))
        
        for i, field in enumerate(fields):
            label_text = field["label"]
            if field["required"]:
                label_text += " *"
                
            label = ctk.CTkLabel(
                fields_frame,
                text=label_text,
                font=ctk.CTkFont(size=13),
                width=120,
                anchor="w"
            )
            label.grid(row=i, column=0, padx=(0, 10), pady=8, sticky="w")
            
            entry = ctk.CTkEntry(fields_frame, height=35, width=400)
            entry.grid(row=i, column=1, pady=8, sticky="ew")
            
            # Store reference to entry
            self.entry_widgets[field["var_name"]] = entry
        
        fields_frame.columnconfigure(1, weight=1)
        return section
    
    def create_fancy_section(self, parent, title, icon):
        section = ctk.CTkFrame(parent, corner_radius=10, border_width=1, border_color=("gray85", "gray40"))
        section.pack(fill="x", pady=10)
        
        # Header
        header = ctk.CTkFrame(section, fg_color=("gray95", "gray25"), corner_radius=8)
        header.pack(fill="x", padx=2, pady=2)
        
        icon_label = ctk.CTkLabel(
            header,
            text=icon,
            font=ctk.CTkFont(size=18),
            width=30
        )
        icon_label.pack(side="left", padx=(15, 5), pady=10)
        
        title_label = ctk.CTkLabel(
            header,
            text=title,
            font=ctk.CTkFont(size=16, weight="bold")
        )
        title_label.pack(side="left", pady=10)
        
        return section
    
    def create_textbox(self, parent, height=100):
        textbox = ctk.CTkTextbox(
            parent, 
            height=height,
            corner_radius=6,
            border_width=1,
            border_color=("gray85", "gray40"),
            wrap="word"
        )
        return textbox
    
    def create_preview_elements(self):
        # Preview header
        header_frame = ctk.CTkFrame(self.preview_frame, fg_color="transparent")
        header_frame.pack(fill="x", padx=20, pady=(20, 15))
        
        header_title = ctk.CTkLabel(
            header_frame,
            text="Resume Preview",
            font=ctk.CTkFont(family="Arial", size=24, weight="bold")
        )
        header_title.pack(side="left")
        
        refresh_button = ctk.CTkButton(
            header_frame,
            text="Refresh Preview",
            font=ctk.CTkFont(size=13),
            width=140,
            height=32,
            corner_radius=6,
            command=self.update_preview
        )
        refresh_button.pack(side="right")
        
        # Preview content
        self.preview_content = ctk.CTkScrollableFrame(self.preview_frame)
        self.preview_content.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        
        self.preview_name = ctk.CTkLabel(
            self.preview_content,
            text="Your Name",
            font=ctk.CTkFont(family="Arial", size=26, weight="bold")
        )
        self.preview_name.pack(anchor="w", pady=(0, 5))
        
        self.preview_contact = ctk.CTkLabel(
            self.preview_content,
            text="email@example.com | 123-456-7890 | Your Address",
            font=ctk.CTkFont(size=14)
        )
        self.preview_contact.pack(anchor="w", pady=(0, 15))
        
        # Summary section
        self.preview_add_section("PROFESSIONAL SUMMARY")
        self.preview_summary = ctk.CTkLabel(
            self.preview_content,
            text="No summary provided yet.",
            font=ctk.CTkFont(size=14),
            wraplength=600,
            justify="left"
        )
        self.preview_summary.pack(anchor="w", pady=(0, 15))
        
        # Skills section
        self.preview_add_section("SKILLS")
        self.preview_skills = ctk.CTkLabel(
            self.preview_content,
            text="No skills provided yet.",
            font=ctk.CTkFont(size=14),
            wraplength=600,
            justify="left"
        )
        self.preview_skills.pack(anchor="w", pady=(0, 15))
        
        # Education section
        self.preview_add_section("EDUCATION")
        self.preview_education = ctk.CTkLabel(
            self.preview_content,
            text="No education provided yet.",
            font=ctk.CTkFont(size=14),
            wraplength=600,
            justify="left"
        )
        self.preview_education.pack(anchor="w", pady=(0, 15))
        
        # Experience section
        self.preview_add_section("EXPERIENCE")
        self.preview_experience = ctk.CTkLabel(
            self.preview_content,
            text="No experience provided yet.",
            font=ctk.CTkFont(size=14),
            wraplength=600,
            justify="left"
        )
        self.preview_experience.pack(anchor="w", pady=(0, 15))
    
    def preview_add_section(self, title):
        section_title = ctk.CTkLabel(
            self.preview_content,
            text=title,
            font=ctk.CTkFont(family="Arial", size=18, weight="bold")
        )
        section_title.pack(anchor="w", pady=(10, 5))
        
        divider = ctk.CTkFrame(self.preview_content, height=2)
        divider.pack(fill="x", pady=(0, 10))
    
    def create_export_elements(self):
        # Export header
        header_frame = ctk.CTkFrame(self.export_frame, fg_color="transparent")
        header_frame.pack(fill="x", padx=20, pady=(20, 15))
        
        header_title = ctk.CTkLabel(
            header_frame,
            text="Export Your Resume",
            font=ctk.CTkFont(family="Arial", size=24, weight="bold")
        )
        header_title.pack(anchor="w")
        
        # Export options container
        export_container = ctk.CTkFrame(self.export_frame, fg_color="transparent")
        export_container.pack(fill="both", expand=True, padx=20, pady=(0, 20))

        # Template selection section
        template_frame = ctk.CTkFrame(export_container, corner_radius=10, border_width=1, border_color=("gray85", "gray40"))
        template_frame.pack(fill="x", pady=10, ipady=15)
        
        template_title = ctk.CTkLabel(
            template_frame,
            text="Select Resume Template",
            font=ctk.CTkFont(family="Arial", size=18, weight="bold"),
        )
        template_title.pack(anchor="w", padx=20, pady=(15, 5))
        
        # Template options
        template_options_frame = ctk.CTkFrame(template_frame, fg_color="transparent")
        template_options_frame.pack(fill="x", padx=20, pady=(5, 15))
        
        self.template_var = ctk.StringVar(value="modern")
        templates = [
            ("Modern", "modern"),
            ("Classic", "classic"),
            ("Minimalist", "minimalist"),
            ("Creative", "creative")
        ]
        
        for i, (text, value) in enumerate(templates):
            template_option = ctk.CTkRadioButton(
                template_options_frame,
                text=text,
                value=value,
                variable=self.template_var,
                font=ctk.CTkFont(size=14)
            )
            template_option.grid(row=0, column=i, padx=15)
        
        # Color scheme section
        color_frame = ctk.CTkFrame(export_container, corner_radius=10, border_width=1, border_color=("gray85", "gray40"))
        color_frame.pack(fill="x", pady=10, ipady=15)
        
        color_title = ctk.CTkLabel(
            color_frame,
            text="Select Color Scheme",
            font=ctk.CTkFont(family="Arial", size=18, weight="bold"),
        )
        color_title.pack(anchor="w", padx=20, pady=(15, 5))
        
        # Color options
        color_options_frame = ctk.CTkFrame(color_frame, fg_color="transparent")
        color_options_frame.pack(fill="x", padx=20, pady=(5, 15))
        
        self.color_var = ctk.StringVar(value="blue")
        colors = [
            ("Blue", "blue"),
            ("Green", "green"),
            ("Purple", "purple"),
            ("Maroon", "maroon"),
            ("Teal", "teal")
        ]
        
        for i, (text, value) in enumerate(colors):
            color_option = ctk.CTkRadioButton(
                color_options_frame,
                text=text,
                value=value,
                variable=self.color_var,
                font=ctk.CTkFont(size=14)
            )
            color_option.grid(row=0, column=i, padx=10)
        
        # PDF option
        pdf_frame = ctk.CTkFrame(export_container, corner_radius=10, border_width=1, border_color=("gray85", "gray40"))
        pdf_frame.pack(fill="x", pady=10, ipady=15)
        
        pdf_title = ctk.CTkLabel(
            pdf_frame,
            text="PDF Format",
            font=ctk.CTkFont(family="Arial", size=18, weight="bold"),
        )
        pdf_title.pack(anchor="w", padx=20, pady=(15, 5))
        
        pdf_desc = ctk.CTkLabel(
            pdf_frame,
            text="Export your resume as a PDF document, ideal for printing and digital submission.",
            font=ctk.CTkFont(size=14),
            wraplength=600,
            justify="left"
        )
        pdf_desc.pack(anchor="w", padx=20, pady=(0, 15))
        
        pdf_button = ctk.CTkButton(
            pdf_frame,
            text="Generate PDF",
            font=ctk.CTkFont(size=14),
            width=160,
            height=40,
            corner_radius=8,
            command=self.generate_pdf
        )
        pdf_button.pack(anchor="w", padx=20)
        
        # Word option
        word_frame = ctk.CTkFrame(export_container, corner_radius=10, border_width=1, border_color=("gray85", "gray40"))
        word_frame.pack(fill="x", pady=10, ipady=15)
        
        word_title = ctk.CTkLabel(
            word_frame,
            text="Word Document",
            font=ctk.CTkFont(family="Arial", size=18, weight="bold"),
        )
        word_title.pack(anchor="w", padx=20, pady=(15, 5))
        
        word_desc = ctk.CTkLabel(
            word_frame,
            text="Export your resume as a Word document, perfect for further editing and customization.",
            font=ctk.CTkFont(size=14),
            wraplength=600,
            justify="left"
        )
        word_desc.pack(anchor="w", padx=20, pady=(0, 15))
        
        word_button = ctk.CTkButton(
            word_frame,
            text="Generate Word Document",
            font=ctk.CTkFont(size=14),
            width=210,
            height=40,
            corner_radius=8,
            command=self.generate_word
        )
        word_button.pack(anchor="w", padx=20)
        
    def generate_pdf(self):
        data = self.get_user_data()
        template = self.template_var.get()
        color_scheme = self.color_var.get()
        
        # Validate data
        if not data['name'] or not data['email']:
            messagebox.showerror("Error", "Name and Email are required fields")
            return
            
        # Create PDF resume
        filename = f"{data['name'].replace(' ', '_')}_Resume.pdf"
        doc = SimpleDocTemplate(filename, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Define colors based on color scheme
        color_map = {
            "blue": colors.navy,
            "green": colors.darkgreen,
            "purple": colors.purple,
            "maroon": colors.maroon,
            "teal": colors.teal
        }
        
        main_color = color_map.get(color_scheme, colors.navy)
        
        # Create custom styles based on template
        if template == "modern":
            title_style = ParagraphStyle(
                'Title',
                parent=styles['Title'],
                fontSize=24,
                alignment=1,
                spaceAfter=20,
                textColor=main_color
            )
            
            heading_style = ParagraphStyle(
                'Heading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=main_color,
                spaceAfter=10,
                borderWidth=0,
                borderPadding=0,
                borderColor=main_color,
                borderRadius=None
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                spaceBefore=6,
                spaceAfter=6
            )
            
        elif template == "classic":
            title_style = ParagraphStyle(
                'Title',
                parent=styles['Title'],
                fontSize=20,
                alignment=1,
                spaceAfter=15,
                textColor=colors.black
            )
            
            heading_style = ParagraphStyle(
                'Heading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=main_color,
                spaceAfter=6,
                borderWidth=0,
                borderPadding=5,
                underline=1
            )
            
            normal_style = styles['Normal']
            
        elif template == "minimalist":
            title_style = ParagraphStyle(
                'Title',
                parent=styles['Title'],
                fontSize=18,
                alignment=0,  # Left aligned
                spaceAfter=15,
                textColor=colors.black
            )
            
            heading_style = ParagraphStyle(
                'Heading',
                parent=styles['Heading2'],
                fontSize=12,
                textColor=main_color,
                spaceAfter=8,
                fontName='Helvetica-Bold'
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                fontName='Helvetica'
            )
            
        else:  # creative
            title_style = ParagraphStyle(
                'Title',
                parent=styles['Title'],
                fontSize=26,
                alignment=1,
                spaceAfter=25,
                textColor=main_color,
                fontName='Helvetica-Bold'
            )
            
            heading_style = ParagraphStyle(
                'Heading',
                parent=styles['Heading2'],
                fontSize=16,
                textColor=main_color,
                spaceAfter=12,
                fontName='Helvetica-Bold'
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceBefore=8,
                spaceAfter=8,
                fontName='Helvetica'
            )
        
        content = []
        
        # Title
        content.append(Paragraph(f"{data['name']}", title_style))
        
        # Contact information
        contact_parts = []
        if data['email']:
            contact_parts.append(f"<b>Email:</b> {data['email']}")
        if data['phone']:
            contact_parts.append(f"<b>Phone:</b> {data['phone']}")
        if data['address']:
            contact_parts.append(f"<b>Address:</b> {data['address']}")
            
        contact_info = " | ".join(contact_parts)
        content.append(Paragraph(contact_info, normal_style))
        content.append(Spacer(1, 15))
        
        # Professional Summary
        content.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
        content.append(Paragraph(data['summary'], normal_style))
        content.append(Spacer(1, 15))
        
        # Skills
        content.append(Paragraph("SKILLS", heading_style))
        skills_list = data['skills'].split(',')
        
        if template == "creative":
            # Creative template uses bullet points for skills
            skills_text = "<ul>"
            for skill in skills_list:
                if skill.strip():
                    skills_text += f"<li>{skill.strip()}</li>"
            skills_text += "</ul>"
        else:
            # Other templates use comma-separated list
            skills_text = ", ".join([skill.strip() for skill in skills_list if skill.strip()])
            
        content.append(Paragraph(skills_text, normal_style))
        content.append(Spacer(1, 15))
        
        # Education
        content.append(Paragraph("EDUCATION", heading_style))
        content.append(Paragraph(data['education'], normal_style))
        content.append(Spacer(1, 15))
        
        # Experience
        content.append(Paragraph("EXPERIENCE", heading_style))
        content.append(Paragraph(data['experience'], normal_style))
        
        # Build document
        doc.build(content)
        
        messagebox.showinfo("Success", f"PDF Resume has been saved as '{filename}'")
        
        # Open the file after generation
        open_generated_file(filename)
    
    def generate_word(self):
        data = self.get_user_data()
        template = self.template_var.get()
        color_scheme = self.color_var.get()
        
        # Validate data
        if not data['name'] or not data['email']:
            messagebox.showerror("Error", "Name and Email are required fields")
            return
            
        # Create Word Document
        doc = Document()
        
        # Color code mapping
        color_map = {
            "blue": "0000FF",
            "green": "008000",
            "purple": "800080",
            "maroon": "800000",
            "teal": "008080"
        }
        
        hex_color = color_map.get(color_scheme, "0000FF")
        
        # Template-specific styling
        if template == "modern":
            # Add title with modern styling
            title = doc.add_heading("", 0)
            title_run = title.add_run(data['name'])
            title_run.font.size = docx.shared.Pt(24)
            title_run.font.color.rgb = docx.shared.RGBColor.from_string(hex_color)
            title.alignment = 1  # Center alignment
            
            # Contact info with horizontal line
            doc.add_paragraph().add_run().add_break()
            contact_info = f"Email: {data['email']} | Phone: {data['phone']} | Address: {data['address']}"
            contact = doc.add_paragraph()
            contact.alignment = 1  # Center alignment
            contact.add_run(contact_info)
            
            # Add horizontal line
            p = doc.add_paragraph()
            p.paragraph_format.space_after = docx.shared.Pt(20)
            p_fmt = p.paragraph_format
            p_fmt.border_bottom = True
            p_fmt.border_bottom_color = hex_color
            p_fmt.border_bottom_width = docx.shared.Pt(1)
            
        elif template == "classic":
            # Classic title (simple and elegant)
            title = doc.add_heading("", 0)
            title_run = title.add_run(data['name'])
            title_run.font.size = docx.shared.Pt(20)
            title.alignment = 1  # Center alignment
            
            # Contact info
            contact_para = doc.add_paragraph()
            contact_para.alignment = 1  # Center alignment
            contact_para.add_run(f"Email: {data['email']} | Phone: {data['phone']} | Address: {data['address']}")
            doc.add_paragraph().add_run().add_break()
            
        elif template == "minimalist":
            # Minimalist title (left-aligned, clean)
            title = doc.add_heading("", 0)
            title_run = title.add_run(data['name'])
            title_run.font.size = docx.shared.Pt(18)
            title.alignment = 0  # Left alignment
            
            # Simple contact info
            doc.add_paragraph(f"Email: {data['email']}")
            doc.add_paragraph(f"Phone: {data['phone']}")
            doc.add_paragraph(f"Address: {data['address']}")
            doc.add_paragraph().add_run().add_break()
            
        else:  # creative
            # Creative bold title with color
            title = doc.add_heading("", 0)
            title_run = title.add_run(data['name'].upper())
            title_run.font.size = docx.shared.Pt(26)
            title_run.font.color.rgb = docx.shared.RGBColor.from_string(hex_color)
            title.alignment = 1  # Center alignment
            
            # Stylized contact info
            contact_para = doc.add_paragraph()
            contact_para.alignment = 1  # Center alignment
            
            if data['email']:
                contact_para.add_run("Email: ").bold = True
                contact_para.add_run(f"{data['email']} ")
            if data['phone']:
                contact_para.add_run("| Phone: ").bold = True
                contact_para.add_run(f"{data['phone']} ")
            if data['address']:
                contact_para.add_run("| Address: ").bold = True
                contact_para.add_run(data['address'])
                
            # Add decorative line
            p = doc.add_paragraph()
            p.paragraph_format.space_after = docx.shared.Pt(20)
            p_fmt = p.paragraph_format
            p_fmt.border_bottom = True
            p_fmt.border_bottom_color = hex_color
            p_fmt.border_bottom_width = docx.shared.Pt(2)
        
        # Professional Summary section
        self.add_section(doc, "PROFESSIONAL SUMMARY", data['summary'], template, hex_color)
        
        # Skills section
        skills_heading = self.add_section_heading(doc, "SKILLS", template, hex_color)
        
        if template == "creative":
            # Creative template uses bullet points for skills
            skills_para = doc.add_paragraph()
            for skill in data['skills'].split(','):
                if skill.strip():
                    bullet_para = doc.add_paragraph(style='ListBullet')
                    bullet_para.add_run(skill.strip())
        else:
            # Other templates use comma-separated list
            skills_para = doc.add_paragraph()
            skills = data['skills'].split(',')
            for i, skill in enumerate(skills):
                if i > 0:
                    skills_para.add_run(', ')
                skills_para.add_run(skill.strip())
        
        # Education section
        self.add_section(doc, "EDUCATION", data['education'], template, hex_color)
        
        # Experience section
        self.add_section(doc, "EXPERIENCE", data['experience'], template, hex_color)
        
        # Save document
        filename = f"{data['name'].replace(' ', '_')}_Resume.docx"
        doc.save(filename)
        
        messagebox.showinfo("Success", f"Word Resume has been saved as '{filename}'")
        
        # Open the file after generation
        open_generated_file(filename)
    
    def open_file(self, filepath):
        """Opens a file with the default application based on the OS"""
        filepath = os.path.abspath(filepath)
        print(f"Opening file: {filepath}")
        
        # Ensure file exists
        if not os.path.exists(filepath):
            messagebox.showerror("Error", f"The file could not be found: {filepath}")
            return
            
        try:
            # Use the utility function if available
            return open_generated_file(filepath)
        except Exception as e:
            print(f"Error in open_file: {e}")
            messagebox.showwarning(
                "Warning", 
                f"Could not open the file automatically.\nThe file has been saved to:\n{filepath}"
            )
    
    def add_section(self, doc, heading_text, content_text, template, hex_color):
        """Helper method to add a section with consistent formatting"""
        self.add_section_heading(doc, heading_text, template, hex_color)
        
        # Add content with appropriate styling
        content_para = doc.add_paragraph()
        content_para.add_run(content_text)
        
        # Add some space after sections
        doc.add_paragraph().add_run().add_break()
        
    def add_section_heading(self, doc, heading_text, template, hex_color):
        """Helper method to add section headings with consistent formatting"""
        if template == "modern":
            heading = doc.add_heading('', level=1)
            heading_run = heading.add_run(heading_text)
            heading_run.font.color.rgb = docx.shared.RGBColor.from_string(hex_color)
            
        elif template == "classic":
            heading = doc.add_heading('', level=1)
            heading_run = heading.add_run(heading_text)
            heading_run.font.color.rgb = docx.shared.RGBColor.from_string(hex_color)
            heading_run.underline = True
            
        elif template == "minimalist":
            heading = doc.add_paragraph()
            heading_run = heading.add_run(heading_text)
            heading_run.bold = True
            heading_run.font.size = docx.shared.Pt(12)
            # Add a thin line after minimalist headings
            p = doc.add_paragraph()
            p.paragraph_format.space_after = docx.shared.Pt(10)
            p_fmt = p.paragraph_format
            p_fmt.border_bottom = True
            p_fmt.border_bottom_width = docx.shared.Pt(0.5)
            
        else:  # creative
            heading = doc.add_heading('', level=1)
            heading_run = heading.add_run(heading_text)
            heading_run.font.color.rgb = docx.shared.RGBColor.from_string(hex_color)
            heading_run.font.size = docx.shared.Pt(16)
            # Bold and capitalized for creative style
            heading_run.bold = True
            heading_text = heading_text.upper()
            
        return heading
    
    def bind_progress_updates(self):
        # Monitor changes in entry widgets
        for entry in self.entry_widgets.values():
            entry.bind("<KeyRelease>", self.update_progress)
        
        # Monitor changes in textboxes
        self.summary_text.bind("<KeyRelease>", self.update_progress)
        self.skills_text.bind("<KeyRelease>", self.update_progress)
        self.education_text.bind("<KeyRelease>", self.update_progress)
        self.experience_text.bind("<KeyRelease>", self.update_progress)
    
    def update_progress(self, event=None):
        # Count filled fields
        filled_count = 0
        total_count = 8  # All fields
        
        # Check entry widgets
        for name, entry in self.entry_widgets.items():
            if entry.get().strip():
                filled_count += 1
        
        # Check textboxes
        if self.summary_text.get("1.0", "end-1c").strip():
            filled_count += 1
        if self.skills_text.get("1.0", "end-1c").strip():
            filled_count += 1
        if self.education_text.get("1.0", "end-1c").strip():
            filled_count += 1
        if self.experience_text.get("1.0", "end-1c").strip():
            filled_count += 1
        
        # Update progress bar
        progress = filled_count / total_count
        self.progress_bar.set(progress)
    
    def show_form(self):
        self.hide_all_frames()
        self.form_frame.pack(fill="both", expand=True)
        self.active_view = "form"
        self.update_nav_highlight()
    
    def show_preview(self):
        self.hide_all_frames()
        self.preview_frame.pack(fill="both", expand=True)
        self.active_view = "preview"
        self.update_nav_highlight()
        self.update_preview()
    
    def show_export(self):
        self.hide_all_frames()
        self.export_frame.pack(fill="both", expand=True)
        self.active_view = "export"
        self.update_nav_highlight()
    
    def show_help(self):
        self.hide_all_frames()
        self.help_frame.pack(fill="both", expand=True)
        self.active_view = "help"
        self.update_nav_highlight()
    
    def hide_all_frames(self):
        self.form_frame.pack_forget()
        self.preview_frame.pack_forget()
        self.export_frame.pack_forget()
        self.help_frame.pack_forget()
    
    def update_preview(self):
        data = self.get_user_data()
        
        # Update name
        self.preview_name.configure(text=data["name"] if data["name"] else "Your Name")
        
        # Update contact info
        contact_parts = []
        if data["email"]:
            contact_parts.append(data["email"])
        if data["phone"]:
            contact_parts.append(data["phone"])
        if data["address"]:
            contact_parts.append(data["address"])
        
        contact_text = " | ".join(contact_parts) if contact_parts else "email@example.com | 123-456-7890 | Your Address"
        self.preview_contact.configure(text=contact_text)
        
        # Update summary
        self.preview_summary.configure(
            text=data["summary"] if data["summary"] else "No summary provided yet."
        )
        
        # Update skills
        self.preview_skills.configure(
            text=data["skills"] if data["skills"] else "No skills provided yet."
        )
        
        # Update education
        self.preview_education.configure(
            text=data["education"] if data["education"] else "No education provided yet."
        )
        
        # Update experience
        self.preview_experience.configure(
            text=data["experience"] if data["experience"] else "No experience provided yet."
        )
    
    def get_user_data(self):
        return {
            "name": self.entry_widgets["name"].get(),
            "email": self.entry_widgets["email"].get(),
            "phone": self.entry_widgets["phone"].get(),
            "address": self.entry_widgets["address"].get(),
            "summary": self.summary_text.get("1.0", "end-1c"),
            "skills": self.skills_text.get("1.0", "end-1c"),
            "education": self.education_text.get("1.0", "end-1c"),
            "experience": self.experience_text.get("1.0", "end-1c")
        }
    
    def clear_form(self):
        # Clear entry widgets
        for entry in self.entry_widgets.values():
            entry.delete(0, 'end')
        
        # Clear textboxes
        self.summary_text.delete("1.0", "end")
        self.skills_text.delete("1.0", "end")
        self.education_text.delete("1.0", "end")
        self.experience_text.delete("1.0", "end")
        
        # Update progress
        self.progress_bar.set(0)
    
    def run(self):
        self.app.mainloop()

    def create_help_elements(self):
        # Help header
        header_frame = ctk.CTkFrame(self.help_frame, fg_color="transparent")
        header_frame.pack(fill="x", padx=20, pady=(20, 15))
        
        header_title = ctk.CTkLabel(
            header_frame,
            text="Help & Tips",
            font=ctk.CTkFont(family="Arial", size=24, weight="bold")
        )
        header_title.pack(anchor="w")
        
        # Help content
        help_container = ctk.CTkScrollableFrame(self.help_frame)
        help_container.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        
        # Tips sections
        self.create_tip_section(
            help_container,
            "Getting Started",
            "1. Fill in your personal information in the 'Create Resume' tab\n"
            "2. Add your professional summary, skills, education, and experience\n"
            "3. Preview your resume in the 'Preview' tab\n"
            "4. Export your resume as PDF or Word document in the 'Export Options' tab"
        )
        
        self.create_tip_section(
            help_container,
            "Writing an Effective Summary",
            "Your professional summary should be a brief 3-5 sentence paragraph that highlights your "
            "most relevant skills, experience, and achievements. Think of it as your elevator pitch."
        )
        
        self.create_tip_section(
            help_container,
            "Listing Skills",
            "• List skills that are relevant to the job you're applying for\n"
            "• Include a mix of technical and soft skills\n"
            "• Be specific (e.g., 'Python' instead of 'Programming')\n"
            "• Separate skills with commas"
        )
        
        self.create_tip_section(
            help_container,
            "Education Section Tips",
            "• List your highest degree first\n"
            "• Include the degree name, institution, location, and graduation date\n"
            "• Add relevant coursework, honors, or academic achievements if applicable"
        )
        n achievements rather than just duties"
        self.create_tip_section(        )
            help_container,
            "Work Experience Tips",    def create_tip_section(self, parent, title, content):
            "• Use the format: Job Title, Company Name, Location, Dates\n"rame(parent, corner_radius=10, border_width=1, border_color=("gray85", "gray40"))
            "• Describe your responsibilities and achievements using action verbs\n"section.pack(fill="x", pady=10, ipady=10)
            "• Quantify your achievements when possible (e.g., 'Increased sales by 20%')\n"
        )Label(
        return sectionon,

# At the end of the file, add this entry point codeold")

if __name__ == "__main__":
    try:
        # Try to initialize and run the application








        messagebox.showerror("Error", f"An error occurred:\n{str(e)}")        print(error_message)        error_message = f"An error occurred:\n{str(e)}\n\n{traceback.format_exc()}"        import traceback        # If an error occurs, show a message box with the error    except Exception as e:        app.run()        app = ResumeBuilder()            section,
            text=content,
            font=ctk.CTkFont(size=14),
            wraplength=600,
            justify="left"
        )
        content_label.pack(anchor="w", padx=20, pady=(0, 15))
        
        return section

# At the end of the file, add this entry point code

if __name__ == "__main__":
    try:
        # Try to initialize and run the application
        app = ResumeBuilder()
        app.run()
    except Exception as e:
        # If an error occurs, show a message box with the error
        import traceback
        error_message = f"An error occurred:\n{str(e)}\n\n{traceback.format_exc()}"
        print(error_message)
        messagebox.showerror("Error", f"An error occurred:\n{str(e)}")